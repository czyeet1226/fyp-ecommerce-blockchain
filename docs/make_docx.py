"""
Renders all activity diagrams to PNG and assembles them into a Word document.

    py make_docx.py
"""

import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

from render_activity import OUT_DIR, render
from specs_activity import DIAGRAMS
from schema_tables_data import SCHEMA_TABLES
import render_wireframe
from specs_wireframe import SCREENS
from specs_testcases import GROUPS as TEST_CASE_GROUPS
from specs_uat import FORMS as UAT_FORMS, FEEDBACK_QUESTIONS

WF_DIR = render_wireframe.OUT_DIR

DOCX_PATH = os.path.join(os.path.dirname(OUT_DIR), "Activity-Diagrams.docx")

# Usable area on a portrait A4/Letter page with 0.6in margins.
MAX_W_IN = 7.0
MAX_H_IN = 9.1


def render_all():
    os.makedirs(OUT_DIR, exist_ok=True)
    paths = []
    for spec in DIAGRAMS:
        p = os.path.join(OUT_DIR, f"{spec['num']}.png")
        render(spec, p)
        paths.append((spec, p))
        print(f"rendered {spec['num']} {spec['title']}")
    return paths


def fitted_size(png):
    """Scale so the image fits the page box without distortion."""
    with Image.open(png) as im:
        w, h = im.size
    # treat the rendered pixels as 96 dpi
    w_in, h_in = w / 96, h / 96
    scale = min(MAX_W_IN / w_in, MAX_H_IN / h_in, 1.0)
    return Inches(w_in * scale), Inches(h_in * scale)


def render_wireframes():
    os.makedirs(WF_DIR, exist_ok=True)
    paths = []
    for spec in SCREENS:
        p = os.path.join(WF_DIR, f"{spec['num']}.png")
        render_wireframe.render(spec, p)
        paths.append((spec, p))
        print(f"wireframe {spec['num']} {spec['title']}")
    return paths


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    tcPr.append(shd)


def add_schema_table(doc, entity):
    heading = doc.add_heading(f"{entity['name']} {entity['store'].split()[-1]}", level=2)
    doc.add_paragraph(entity["store"]).runs[0].font.italic = True

    if entity.get("description"):
        desc = doc.add_paragraph(entity["description"])
        desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(0.85), Inches(1.7), Inches(1.55), Inches(2.9)]
    headers = ["Data Type", "Field Name", "Constraints", "Description"]

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].width = widths[i]
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], "D9D9D9")

    for data_type, field_name, constraints, description in entity["fields"]:
        row_cells = table.add_row().cells
        values = [data_type, field_name, constraints, description]
        for i, val in enumerate(values):
            row_cells[i].text = val
            row_cells[i].width = widths[i]

    doc.add_paragraph()


def add_test_case_table(doc, group):
    doc.add_heading(f"{group['usecase']}", level=2)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(0.55), Inches(1.75), Inches(1.75), Inches(1.75),
              Inches(0.85), Inches(0.65)]
    headers = ["TCNO", "User Input", "System Action", "Expected Output",
               "Actual Output", "Result"]

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].width = widths[i]
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], "D9D9D9")

    for tcno, user_input, system_action, expected in group["cases"]:
        row_cells = table.add_row().cells
        values = [tcno, user_input, system_action, expected, "", ""]
        for i, val in enumerate(values):
            row_cells[i].text = val
            row_cells[i].width = widths[i]

    cap = doc.add_paragraph(
        f"Table {group['table_no']}: {group['title']}"
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def add_test_case_section(doc):
    doc.add_page_break()
    h = doc.add_heading("Test Cases", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Blockchain E-Commerce Platform \u2014 Customer, Merchant and Admin "
        "Test Case Tables"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Each table follows the TCNO / User Input / System Action / Expected "
        "Output / Actual Output / Result format, with Actual Output and "
        "Result left blank for the tester to complete during execution. Test "
        "cases are grouped by use case (see docs/usecase-specifications.md) so "
        "each row traces back to a specific Basic, Alternative, or Exception "
        "flow and the backend route it exercises."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    doc.add_paragraph("Contents").runs[0].bold = True
    last_actor = None
    for group in TEST_CASE_GROUPS:
        if group["actor"] != last_actor:
            gp = doc.add_paragraph(group["actor"])
            gp.runs[0].bold = True
            gp.paragraph_format.space_before = Pt(6)
            last_actor = group["actor"]
        doc.add_paragraph(f"Table {group['table_no']}: {group['title']}",
                          style="List Bullet")

    last_actor = None
    for group in TEST_CASE_GROUPS:
        doc.add_page_break()
        if group["actor"] != last_actor:
            doc.add_heading(group["actor"], level=1)
            last_actor = group["actor"]
        add_test_case_table(doc, group)


def add_uat_form(doc, form):
    doc.add_heading(f"{form['section_no']}. {form['title']}", level=2)

    n_cols = 5
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(2.4), Inches(1.15), Inches(1.15), Inches(1.15),
              Inches(1.15)]

    def merged_row(text, bold=True):
        row = table.add_row()
        cell = row.cells[0]
        for c in row.cells[1:]:
            cell = cell.merge(c)
        cell.text = text
        cell.paragraphs[0].runs[0].bold = bold
        return cell

    # tester info block (each line spans the full table width)
    merged_row("Name of the Tester:")
    merged_row("Job Position:")
    merged_row("Date:")
    merged_row("Time Start:                                    Time End:")

    # rating grid header
    header_row = table.add_row()
    headers = ["Name of parameter/Testing Category", "Excellent\n(5) points",
               "Very Good\n(4) points", "Fair\n(2-3) points",
               "Poor (1-0)\npoints"]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.width = widths[i]
        p = cell.paragraphs[0]
        lines = text.split("\n")
        p.text = lines[0]
        p.runs[0].bold = True
        for extra in lines[1:]:
            p2 = cell.add_paragraph(extra)
            p2.runs[0].bold = True if p2.runs else None
        set_cell_shading(cell, "D9D9D9")

    # one row per testing category, rating cells left blank for the tester
    for category in form["categories"]:
        row = table.add_row()
        row.cells[0].text = category
        for i in range(1, n_cols):
            row.cells[i].text = ""
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # open feedback questions + developer response, each spanning full width
    for i, q in enumerate(FEEDBACK_QUESTIONS, start=1):
        cell = merged_row(f"{i}. {q}")
        cell.add_paragraph(":")

    resp_cell = merged_row("Developer Response")
    resp_cell.add_paragraph(":")

    doc.add_paragraph()


def add_uat_section(doc):
    doc.add_page_break()
    h = doc.add_heading("UAT Preparation", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Blockchain E-Commerce Platform \u2014 User Acceptance Testing Forms "
        "by Role"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Each form lets a tester score every major feature area for a given "
        "role on a 0-5 scale (Excellent 5, Very Good 4, Fair 2-3, Poor 1-0), "
        "followed by two open feedback questions and a space for the "
        "developer's response. Testing categories are derived from the "
        "screens and use cases documented earlier in this file "
        "(usecase-specifications.md and the UI wireframes), so each row "
        "corresponds to a real, testable feature of the platform."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    doc.add_paragraph("Contents").runs[0].bold = True
    for form in UAT_FORMS:
        doc.add_paragraph(f"{form['section_no']}.  {form['title']}",
                          style="List Bullet")

    for form in UAT_FORMS:
        doc.add_page_break()
        add_uat_form(doc, form)


def build(paths, wireframes):
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(0.6))

    h = doc.add_heading("Activity Diagrams", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Blockchain E-Commerce Platform")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Each diagram separates actor actions from system and blockchain "
        "responsibilities using two swimlanes. A solid filled circle marks the "
        "start node and a ringed circle with a filled centre marks the end "
        "node. Diamonds are decision points and every branch is labelled."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    doc.add_paragraph("Contents").runs[0].bold = True
    for spec in DIAGRAMS:
        doc.add_paragraph(f"{int(spec['num'])}.  {spec['title']}",
                          style="List Bullet")

    for spec, png in paths:
        doc.add_page_break()
        doc.add_heading(f"{int(spec['num'])}. {spec['title']}", level=1)
        lanes = "  |  ".join(spec["lanes"])
        cap = doc.add_paragraph(f"Swimlanes: {lanes}")
        cap.runs[0].font.size = Pt(10)
        cap.runs[0].font.italic = True

        w, hh = fitted_size(png)
        doc.add_picture(png, width=w, height=hh)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    h2 = doc.add_heading("Database Schema Tables", level=0)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(
        "Blockchain E-Commerce Platform \u2014 MySQL Tables and MongoDB Collection"
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14)
    p2.runs[0].font.italic = True

    doc.add_paragraph()
    intro2 = doc.add_paragraph(
        "Field-level layout of every table generated from the actual schema in "
        "backend/models/mysql.models.js (MySQL) and "
        "backend/models/blockchainLog.model.js (MongoDB). Each table lists the "
        "data type, field name, constraints, and description for every column."
    )
    intro2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for entity in SCHEMA_TABLES:
        doc.add_page_break()
        add_schema_table(doc, entity)

    add_wireframe_section(doc, wireframes)
    add_test_case_section(doc)
    add_uat_section(doc)

    doc.save(DOCX_PATH)
    print(f"\nsaved {DOCX_PATH}")


def add_wireframe_section(doc, wireframes):
    doc.add_page_break()
    h = doc.add_heading("User Interface Wireframes", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Blockchain E-Commerce Platform \u2014 Customer, Seller and Admin Screens"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Low-fidelity wireframes for every screen in the system, laid out from "
        "the actual React implementation so navigation labels, form fields, "
        "table columns, button text, staking tiers, payment modes and plan "
        "limits match the running application. Conditional states that change "
        "what the user can do \u2014 such as MetaMask connected versus "
        "disconnected, or the escrow lifecycle \u2014 are wireframed "
        "separately. Dashed callouts explain behaviour that is not visible from "
        "the static layout alone."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    doc.add_paragraph("Contents").runs[0].bold = True
    last_group = None
    for spec in SCREENS:
        if spec["group"] != last_group:
            gp = doc.add_paragraph(spec["group"])
            gp.runs[0].bold = True
            gp.paragraph_format.space_before = Pt(6)
            last_group = spec["group"]
        doc.add_paragraph(f"{int(spec['num'])}.  {spec['title']}",
                          style="List Bullet")

    last_group = None
    for spec, png in wireframes:
        doc.add_page_break()
        if spec["group"] != last_group:
            doc.add_heading(spec["group"], level=1)
            last_group = spec["group"]
        doc.add_heading(f"{int(spec['num'])}. {spec['title']}", level=2)
        cap = doc.add_paragraph(spec["caption"])
        cap.runs[0].font.size = Pt(10)
        cap.runs[0].font.italic = True

        w, hh = fitted_size(png)
        doc.add_picture(png, width=w, height=hh)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    build(render_all(), render_wireframes())
